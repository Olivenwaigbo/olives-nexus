# file_loader.py
# Multi-format ingestion for Olive's Nexus.
# Turns CSV / Excel / PDF / Word / plain text into:
#   - a pandas DataFrame, when tabular data can be found, and/or
#   - extracted free text, for documents that are mostly narrative.
#
# app.py decides what to do when df is None (no table found) — usually:
# fall back to extract_table_from_text() in llm.py, or just use the text as
# extra context for the dashboard prompt.

import io
import pandas as pd

SUPPORTED_TYPES = ["csv", "xlsx", "xls", "pdf", "docx", "txt", "md"]


# ── Individual format loaders ───────────────────────────────────────────────
def load_csv(file) -> pd.DataFrame:
    file.seek(0)
    return pd.read_csv(file)


def load_excel(file) -> pd.DataFrame:
    file.seek(0)
    return pd.read_excel(file, engine="openpyxl")


def _coerce_numeric(df: pd.DataFrame) -> pd.DataFrame:
    """Try to convert object columns that are actually numeric."""
    for col in df.columns:
        converted = pd.to_numeric(
            df[col].astype(str).str.replace(",", "").str.replace("$", "", regex=False),
            errors="coerce",
        )
        # Only swap in the numeric version if most values actually converted
        if converted.notna().sum() >= max(1, int(len(df) * 0.6)):
            df[col] = converted
    return df


def load_pdf(file):
    """
    Returns (df_or_None, full_text).
    Pulls the largest embedded table (if any) plus all page text.
    """
    import pdfplumber

    file.seek(0)
    text_parts, tables = [], []
    with pdfplumber.open(file) as pdf:
        for page in pdf.pages:
            text_parts.append(page.extract_text() or "")
            for table in page.extract_tables():
                if table and len(table) > 1 and len(table[0]) > 1:
                    tables.append(table)

    full_text = "\n".join(text_parts).strip()
    df = None
    if tables:
        best = max(tables, key=lambda t: len(t) * len(t[0]))
        header, *rows = best
        header = [str(h).strip() if h else f"col_{i}" for i, h in enumerate(header)]
        # pad/truncate rows to header length defensively
        rows = [r + [None] * (len(header) - len(r)) for r in rows]
        df = pd.DataFrame(rows, columns=header)
        df = _coerce_numeric(df)

    return df, (full_text or None)


def load_docx(file):
    """Returns (df_or_None, full_text). Pulls the largest table + all paragraph text."""
    from docx import Document

    file.seek(0)
    doc = Document(file)
    text_parts = [p.text for p in doc.paragraphs if p.text.strip()]
    full_text = "\n".join(text_parts).strip()

    df = None
    if doc.tables:
        best = max(doc.tables, key=lambda t: len(t.rows) * len(t.columns))
        if len(best.rows) > 1:
            header = [c.text.strip() for c in best.rows[0].cells]
            rows = [[c.text.strip() for c in r.cells] for r in best.rows[1:]]
            df = pd.DataFrame(rows, columns=header)
            df = _coerce_numeric(df)

    return df, (full_text or None)


def load_txt(file):
    """Returns (df_or_None, raw_text). Tries to sniff a delimited table inside plain text."""
    file.seek(0)
    raw = file.read()
    if isinstance(raw, bytes):
        raw = raw.decode("utf-8", errors="ignore")

    df = None
    try:
        candidate = pd.read_csv(io.StringIO(raw), sep=None, engine="python")
        if candidate.shape[1] > 1 and candidate.shape[0] > 0:
            df = _coerce_numeric(candidate)
    except Exception:
        pass

    return df, raw


# ── Unified entry point ─────────────────────────────────────────────────────
def load_any(uploaded_file):
    """
    Dispatch based on file extension.

    Returns dict:
      {
        "df":          DataFrame or None,
        "text":        extracted text or None (None for clean csv/xlsx),
        "source_type": "csv" | "xlsx" | "xls" | "pdf" | "docx" | "txt" | "md",
        "filename":    original filename,
      }
    """
    name = uploaded_file.name
    ext = name.rsplit(".", 1)[-1].lower() if "." in name else ""

    if ext not in SUPPORTED_TYPES:
        raise ValueError(
            f"Unsupported file type: .{ext}. Supported: {', '.join(SUPPORTED_TYPES)}"
        )

    df, text = None, None

    if ext == "csv":
        df = load_csv(uploaded_file)
    elif ext in ("xlsx", "xls"):
        df = load_excel(uploaded_file)
    elif ext == "pdf":
        df, text = load_pdf(uploaded_file)
    elif ext == "docx":
        df, text = load_docx(uploaded_file)
    elif ext in ("txt", "md"):
        df, text = load_txt(uploaded_file)

    return {"df": df, "text": text, "source_type": ext, "filename": name}
